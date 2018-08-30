# -*- coding: utf-8 -*-
import os, sys
from docx import Document

reload(sys)
sys.setdefaultencoding("utf-8")

sys.path.append(os.path.join(os.path.split(os.path.realpath(__file__))[0], '../util'))
import loghelper, db

#logger
loghelper.init_logger("gen_quarterly_report", stream=True)
logger = loghelper.get_logger("gen_quarterly_report")


def get_0_title_paragraphs():
    f = open("samples/QReport_0_title.docx", "rb")
    document = Document(f)
    f.close()
    return document.paragraphs

def get_sample_document():
    f = open("samples/QReport.docx", "rb")
    document = Document(f)
    f.close()
    return document


def main():
    document = get_sample_document()
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            print run.text
            if run.text == "title":
                run.text = u"用户定义的名字"

    for table in document.tables:
        print table.style.name
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text == "title":
                            run.text = u"用户定义的名字"
                        elif run.text == "fullname":
                            run.text = u"烯牛数据测试"
                        elif run.text == "description":
                            run.text = u"烯牛数据描述测试"
                        elif run.text == "table1":
                            print run.text
                            row1 = table.add_row()


    document.save("result/test.docx")


if __name__ == "__main__":
    main()

